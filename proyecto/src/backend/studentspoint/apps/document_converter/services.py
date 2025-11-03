"""
Servicios de conversión de documentos
"""
import os
import logging
from pathlib import Path
from django.core.files.base import ContentFile
from django.utils import timezone
from .utils import DocumentValidator, FileValidationError

logger = logging.getLogger(__name__)


class DocumentConverter:
    """Conversor de documentos Word-PDF y PDF-Word"""
    
    @staticmethod
    def word_to_pdf(word_file_path, output_path=None):
        """Convierte Word a PDF usando python-docx y reportlab
        
        Args:
            word_file_path: Ruta al archivo Word
            output_path: Ruta de salida para el PDF (opcional)
            
        Returns:
            str: Ruta al archivo PDF generado
            
        Raises:
            FileValidationError: Si el archivo es inválido
            Exception: Si hay error en la conversión
        """
        # Validar que el archivo existe y es válido
        is_valid, error_msg = DocumentValidator.validate_file_path(word_file_path)
        if not is_valid:
            raise FileValidationError(error_msg)
        
        try:
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
            from reportlab.lib.utils import ImageReader
            
            # Leer documento Word
            try:
                doc = Document(word_file_path)
            except Exception as e:
                raise FileValidationError(f"El archivo Word está corrupto o no es válido: {str(e)}")
            
            # Verificar que el documento tenga contenido
            if not doc.paragraphs:
                logger.warning(f"Documento Word vacío: {word_file_path}")
                # Crear un PDF con mensaje de que el documento estaba vacío
                pass
            
            # Crear PDF
            if output_path is None:
                base_path = Path(word_file_path)
                output_path = str(base_path.with_suffix('.pdf'))
            
            # Sanitizar nombre de archivo de salida (solo el nombre, no la ruta completa)
            output_path_obj = Path(output_path)
            sanitized_name = DocumentValidator.sanitize_filename(output_path_obj.name)
            output_path = str(output_path_obj.parent / sanitized_name)
            
            pdf = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Estilos personalizados
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor='#1a0933',
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                alignment=TA_JUSTIFY,
                spaceAfter=12
            )
            
            # Procesar párrafos del Word
            has_content = False
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    has_content = True
                    # Detectar si es título
                    if i == 0 or paragraph.style.name.startswith('Heading'):
                        try:
                            story.append(Paragraph(paragraph.text, title_style))
                        except Exception as e:
                            logger.warning(f"Error procesando párrafo {i} como título: {e}")
                            story.append(Paragraph(paragraph.text, normal_style))
                    else:
                        try:
                            story.append(Paragraph(paragraph.text, normal_style))
                        except Exception as e:
                            logger.warning(f"Error procesando párrafo {i}: {e}")
                            # Intentar sin formato especial si hay error
                            story.append(Paragraph(paragraph.text.replace('<', '&lt;').replace('>', '&gt;'), normal_style))
                    story.append(Spacer(1, 0.2 * inch))
            
            # Si no hay contenido, agregar mensaje
            if not has_content:
                story.append(Paragraph("Documento sin contenido", title_style))
            
            # Construir PDF
            try:
                pdf.build(story)
            except Exception as e:
                raise Exception(f"Error generando PDF: {str(e)}")
            
            # Verificar que el PDF se creó correctamente
            if not os.path.exists(output_path):
                raise Exception("El archivo PDF no se generó correctamente")
            
            logger.info(f"Conversion Word a PDF exitosa: {output_path}")
            return output_path
            
        except FileValidationError:
            raise
        except ImportError as e:
            error_msg = f"Librería requerida no disponible: {str(e)}"
            logger.error(error_msg, exc_info=True)
            raise Exception(error_msg)
        except Exception as e:
            error_msg = f"Error convirtiendo Word a PDF: {str(e)}"
            logger.error(error_msg, exc_info=True)
            raise Exception(error_msg)
    
    @staticmethod
    def pdf_to_word(pdf_file_path, output_path=None, usar_ocr=False):
        """Convierte PDF a Word usando PyPDF2 y python-docx"""
        try:
            from PyPDF2 import PdfReader
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Leer PDF
            reader = PdfReader(pdf_file_path)
            
            # Crear documento Word
            doc = Document()
            
            # Configurar estilos
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            
            # Extraer texto de cada página
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                
                if usar_ocr and not text.strip():
                    # Si no hay texto, intentar OCR
                    text = DocumentConverter._extract_with_ocr(pdf_file_path, page_num)
                
                if text.strip():
                    # Agregar número de página
                    heading = doc.add_heading(f'Pagina {page_num + 1}', level=2)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Agregar contenido
                    paragraphs = text.split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            p = doc.add_paragraph(para_text)
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # Separador de página
                    if page_num < len(reader.pages) - 1:
                        doc.add_page_break()
            
            # Guardar documento Word
            if output_path is None:
                output_path = str(pdf_file_path).replace('.pdf', '.docx')
            
            doc.save(output_path)
            
            logger.info(f"Conversion PDF a Word exitosa: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error convirtiendo PDF a Word: {e}", exc_info=True)
            raise
    
    @staticmethod
    def _extract_with_ocr(pdf_path, page_num):
        """Extrae texto usando OCR (pytesseract)"""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            # Convertir página a imagen
            images = convert_from_path(pdf_path, first_page=page_num+1, last_page=page_num+1)
            
            if images:
                # Aplicar OCR
                text = pytesseract.image_to_string(images[0], lang='spa')
                return text
            return ""
            
        except Exception as e:
            logger.warning(f"OCR fallido en pagina {page_num}: {e}")
            return ""


def convert_document(conversion_job):
    """Ejecuta la conversión de documento con manejo robusto de errores
    
    Args:
        conversion_job: Instancia de ConversionJob
        
    Returns:
        bool: True si la conversión fue exitosa, False en caso contrario
    """
    from .models import ConversionJob
    
    output_path = None
    
    try:
        # Validar que el trabajo existe y tiene archivo
        if not conversion_job:
            logger.error("ConversionJob es None")
            return False
        
        if not conversion_job.archivo_original:
            error_msg = "No se proporcionó archivo original"
            conversion_job.estado = ConversionJob.Estado.ERROR
            conversion_job.error_mensaje = error_msg
            conversion_job.save()
            logger.error(f"Error en conversion {conversion_job.id}: {error_msg}")
            return False
        
        # Actualizar estado a procesando
        conversion_job.estado = ConversionJob.Estado.PROCESANDO
        conversion_job.save()
        
        # Obtener rutas
        try:
            input_path = conversion_job.archivo_original.path
        except Exception as e:
            error_msg = f"Error obteniendo ruta del archivo: {str(e)}"
            conversion_job.estado = ConversionJob.Estado.ERROR
            conversion_job.error_mensaje = error_msg
            conversion_job.save()
            logger.error(f"Error en conversion {conversion_job.id}: {error_msg}")
            return False
        
        # Validar ruta de entrada
        is_valid, error_msg = DocumentValidator.validate_file_path(input_path)
        if not is_valid:
            conversion_job.estado = ConversionJob.Estado.ERROR
            conversion_job.error_mensaje = error_msg
            conversion_job.save()
            logger.error(f"Error en conversion {conversion_job.id}: {error_msg}")
            return False
        
        # Ejecutar conversión según tipo
        try:
            if conversion_job.tipo_conversion == ConversionJob.TipoConversion.WORD_TO_PDF:
                output_path = DocumentConverter.word_to_pdf(input_path)
                output_filename = DocumentValidator.sanitize_filename(Path(input_path).stem + '.pdf')
            else:
                output_path = DocumentConverter.pdf_to_word(
                    input_path, 
                    usar_ocr=conversion_job.usar_ocr
                )
                output_filename = DocumentValidator.sanitize_filename(Path(input_path).stem + '.docx')
        except FileValidationError as e:
            error_msg = f"Error de validación: {str(e)}"
            conversion_job.estado = ConversionJob.Estado.ERROR
            conversion_job.error_mensaje = error_msg
            conversion_job.save()
            logger.error(f"Error en conversion {conversion_job.id}: {error_msg}")
            return False
        except Exception as e:
            error_msg = f"Error en la conversión: {str(e)}"
            conversion_job.estado = ConversionJob.Estado.ERROR
            conversion_job.error_mensaje = error_msg
            conversion_job.save()
            logger.error(f"Error en conversion {conversion_job.id}: {error_msg}", exc_info=True)
            return False
        
        # Validar que se generó el archivo de salida
        if not output_path or not os.path.exists(output_path):
            error_msg = "El archivo convertido no se generó correctamente"
            conversion_job.estado = ConversionJob.Estado.ERROR
            conversion_job.error_mensaje = error_msg
            conversion_job.save()
            logger.error(f"Error en conversion {conversion_job.id}: {error_msg}")
            return False
        
        # Guardar archivo convertido
        try:
            with open(output_path, 'rb') as f:
                content = f.read()
                if not content:
                    raise Exception("El archivo convertido está vacío")
                
                conversion_job.archivo_convertido.save(
                    output_filename,
                    ContentFile(content),
                    save=False
                )
        except Exception as e:
            error_msg = f"Error guardando archivo convertido: {str(e)}"
            conversion_job.estado = ConversionJob.Estado.ERROR
            conversion_job.error_mensaje = error_msg
            conversion_job.save()
            logger.error(f"Error en conversion {conversion_job.id}: {error_msg}", exc_info=True)
            # Limpiar archivo temporal si existe
            if output_path and os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except:
                    pass
            return False
        
        # Actualizar estado
        conversion_job.estado = ConversionJob.Estado.COMPLETADO
        conversion_job.completed_at = timezone.now()
        conversion_job.save()
        
        # Limpiar archivo temporal
        if output_path and os.path.exists(output_path):
            try:
                # Solo eliminar si no es el mismo archivo guardado en el modelo
                saved_path = None
                if conversion_job.archivo_convertido:
                    try:
                        saved_path = conversion_job.archivo_convertido.path
                    except:
                        pass
                
                if output_path != saved_path:
                    os.remove(output_path)
            except Exception as e:
                logger.warning(f"No se pudo eliminar archivo temporal {output_path}: {e}")
        
        logger.info(f"Conversion completada exitosamente: {conversion_job.id}")
        return True
        
    except Exception as e:
        error_msg = f"Error inesperado: {str(e)}"
        conversion_job.estado = ConversionJob.Estado.ERROR
        conversion_job.error_mensaje = error_msg
        conversion_job.save()
        logger.error(f"Error inesperado en conversion {conversion_job.id}: {error_msg}", exc_info=True)
        
        # Limpiar archivo temporal si existe
        if output_path and os.path.exists(output_path):
            try:
                os.remove(output_path)
            except:
                pass
        
        return False

